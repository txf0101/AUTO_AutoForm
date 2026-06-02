"""这个工具脚本用于从当前项目资料生成辅助报告或参考资料。运行前应先确认输入路径和输出路径，避免覆盖人工整理的文件。

This utility script generates helper reports or reference material from the current project data. Check input and output paths before running it so manually curated files are not overwritten.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


WORKSPACE = Path(__file__).resolve().parents[1]
OUTPUT_PATH = WORKSPACE / "output" / "doc" / "AutoForm_MCP项目状态与1.0发布差距汇报_20260525.docx"


def main() -> None:
    """Create the Word report requested for the AutoForm MCP status briefing."""
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)

    document = Document()
    configure_document(document)
    add_title(document)
    add_basis_section(document)
    add_command_scope_section(document)
    add_example_project_section(document)
    add_open_source_comparison_section(document)
    add_next_work_section(document)
    add_version_distance_section(document)
    add_appendix_section(document)

    document.save(OUTPUT_PATH)
    print(OUTPUT_PATH)


def configure_document(document: Document) -> None:
    """Apply a compact landscape layout so wide evidence tables remain readable."""
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Cm(1.35)
    section.bottom_margin = Cm(1.35)
    section.left_margin = Cm(1.25)
    section.right_margin = Cm(1.25)

    for style_name, size, bold in [
        ("Normal", 9, False),
        ("Title", 18, True),
        ("Heading 1", 13, True),
        ("Heading 2", 11, True),
    ]:
        style = document.styles[style_name]
        style.font.name = "宋体"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        style.font.size = Pt(size)
        style.font.bold = bold


def add_title(document: Document) -> None:
    """Add the report title and a short scope statement."""
    title = document.add_heading("AutoForm MCP 项目状态与 1.0 发布差距汇报", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("面向当前 AutoForm Agent 工作区与本机 AutoForm Forming R13 环境")
    run.bold = True

    add_paragraph(
        document,
        "本文重新梳理当前 MCP 项目的可接收指令范围、官方示例工程的自然语言到脚本再到软件操作链路、"
        "P0、P1、P2 收口结果，以及相对于公开 Abaqus MCP 项目的工程边界。旧文档结构作为历史材料留存，本报告按 1.0 发布口径重新组织。",
    )
    add_paragraph(
        document,
        "报告日期：2026-05-25。报告中的 AutoForm 事实优先来自本机安装目录、ProgramData、当前工作区源码、"
        "MCP 工具返回值和测试输出。公开项目对比来自 Cai-aa/abaqus-mcp 的 GitHub README 与 GitHub API 元数据。",
    )


def add_basis_section(document: Document) -> None:
    """Record concrete evidence so the briefing can be traced back to files or commands."""
    document.add_heading("一、资料来源与判定边界", level=1)
    add_paragraph(
        document,
        "本次汇报采用证据优先的口径。凡涉及软件路径、版本、工具数量、测试结果和外部项目能力，均在下表列出依据。"
        "对于尚未完成真实执行或缺少官方接口说明的内容，本文使用“待验证”或“工程估计”标注。",
    )

    rows = [
        [
            "B1",
            "本机 AutoForm 安装发现",
            "MCP 工具 autoform_discover_installation 返回 AutoForm Forming R13，版本 13.0.1.02，安装日期 20260519，安装目录为 D:\\Program Files\\AutoForm\\AFplus\\R13F。",
        ],
        [
            "B2",
            "本机 AutoForm 示例工程",
            "MCP 工具 autoform_list_example_projects 返回 7 个官方 .afd 示例：AutoComp_R13、PhaseChange_R13、Sigma_R13、Solver_R13、Thermo_R13、Triboform_R13、Trim_R13。",
        ],
        [
            "B3",
            "当前 MCP 工具入口",
            "autoform_agent\\mcp_server.py 中共有 88 个 @mcp.tool() 入口和 1 个 @mcp.resource() 入口，入口覆盖状态、安装、工程、QuickLink、材料、队列、求解器、作业生命周期、结果证据、发布检查、安全扫描、扩展边界、报告、AF_API 和帮助主题。",
        ],
        [
            "B4",
            "当前 CLI 入口",
            "autoform_agent\\cli.py 中共有 91 个 subparsers.add_parser 入口，CLI 与 MCP 共享底层模块，便于在自然语言工具调用之外进行复核。",
        ],
        [
            "B5",
            "当前测试结果",
            "在 TEMP 与 TMP 指向工作区 tmp\\pytest_runtime_final_<timestamp>，并显式设置 --basetemp tmp\\pytest_basetemp_final_<timestamp> 后执行 <python> -m pytest -q，最终结果为 81 passed in 2.81s。",
        ],
        [
            "B6",
            "模块覆盖矩阵",
            "autoform_agent\\coverage.py 与 autoform_module_coverage_matrix 返回当前已实现模块，其中 Diagnostics 已纳入 autoform_status_snapshot，Simulation jobs 已纳入生命周期工具，Reports 已纳入结果证据包，Release packaging 已纳入发布检查。",
        ],
        [
            "B7",
            "示例工程摘要",
            "autoform_get_afd_project_summary 对 7 个官方示例输出项目名、材料、厚度、DieFace 使用状态和插件使用候选字段；该函数明确说明字段来自 .afd 可读片段，需要用 QuickLink 或官方导出交叉验证。",
        ],
        [
            "B8",
            "公开参照项目",
            "Cai-aa/abaqus-mcp README 说明其 v4.0 支持 Abaqus/CAE 内部插件、文件式 IPC、execute_script、get_model_info、list_jobs、submit_job、get_odb_info、get_viewport_image 和 abaqus://status 资源。",
        ],
        [
            "B9",
            "公开项目元数据",
            "GitHub API 于本轮返回 Cai-aa/abaqus-mcp 为公开 Python 项目，MIT License，默认分支 main，141 stars，17 forks，仓库更新时间 2026-05-24T17:12:38Z。",
        ],
        [
            "B10",
            "V1.0 收口验证",
            "release-readiness 返回 ready=true；public-release-scan 返回 safe_to_publish=true；通过 MCP stdio 客户端调用 autoform_project_run 执行 Solver_R13、Trim_R13 和 AutoComp_R13 三个官方示例，均返回 status=completed、returncode=0、simulation_successful=true，并写出 result_package。",
        ],
        [
            "B11",
            "P1 与 P2 收口证据",
            "docs\\example_project_baselines.json 已记录 7 个官方示例工程基准；quicklink-schema 可输出 schema_version=1.0 的 QuickLink 结构；write-safety-plan 可为 ProgramData 写入生成回滚计划；extension-boundary 已列出外部 CLI、QuickLink Export 脚本和报告模板三类已确认扩展路径。",
        ],
    ]
    add_table(document, ["编号", "依据类别", "具体依据"], rows, [1.6, 5.0, 19.8])


def add_command_scope_section(document: Document) -> None:
    """Describe what instructions the MCP project can currently accept."""
    document.add_heading("二、当前可接收指令范围", level=1)
    add_paragraph(
        document,
        "当前项目已经从“本机 AutoForm 事实读取”扩展到“命令计划、受控执行、日志和结果线索解析”。"
        "从面向汇报的角度，可把能力分为状态、环境、工程、仿真、数据、治理和发布辅助七类。"
    )

    document.add_heading("2.1 总览表", level=2)
    overview_rows = [
        [
            "状态资源与健康检查",
            "读取 autoform://status 或调用 status 命令，汇总项目版本、默认端口、安装记录、队列进程、QuickLink 导出、最近日志、覆盖矩阵和探测错误。",
            "P0 已完成只读入口",
            "B3、B4、B5、B10、diagnostics.py",
        ],
        [
            "环境与安装发现",
            "查找本机 AutoForm 版本、安装目录、bin、材料库、脚本目录、帮助链接和 package_info。",
            "已实现，只读稳定",
            "B1、B3、B6",
        ],
        [
            "工程文件识别",
            "列出官方示例工程，读取 .afd 文件事实，抽取可读片段，形成候选工程摘要，并通过 project-run 生成可复现运行目录。",
            "P1 已完成工程级运行入口",
            "B2、B7、B10、project_workflow.py",
        ],
        [
            "图形界面与工程打开",
            "生成或执行 AutoForm Forming 启动命令，生成或执行 AFFormingUI.exe -file 工程打开命令。",
            "已实现预演与受控执行入口",
            "README、process.py、mcp_server.py",
        ],
        [
            "求解器与作业",
            "生成 AFFormingJob 检查命令，生成 AFFormingSolver 运动学检查与默认完整求解命令，支持批量探测、日志摘要、作业登记、状态刷新、等待、取消和归档计划。",
            "P0 已完成生命周期基础闭环",
            "B6、B10、solver.py、jobs.py、tests\\test_solver.py、tests\\test_jobs.py",
        ],
        [
            "QuickLink 数据桥接",
            "安装 QuickLink 桥接脚本，收集 AutoForm 传出的 QuickLink archive，解析 XML、ProcessPlan、Evaluation、DieFace 和几何引用，并输出 1.0 规范化 schema。",
            "P1 已完成桥接、解析和 schema 化",
            "quicklink.py、quicklink_bridge.py、B6、B11",
        ],
        [
            "材料库治理",
            "安装材料、列出材料库、查找重复材料、生成备份计划、检查 .mat 和 .mtb、执行 .mat 到 .mtb 转换。",
            "已实现，写入动作默认预演",
            "materials.py、commands.py、tests\\test_materials.py",
        ],
        [
            "队列与远程计算",
            "读取队列配置、远程主机配置、日志配置，检查队列进程，生成 AFQueueClient 与 LSF wrapper 命令计划。",
            "已实现只读和探测级能力",
            "config.py、queue.py、B6",
        ],
        [
            "诊断与日志",
            "收集近期日志，解析 GUI 工程打开事件，生成诊断包计划，输出环境快照，并为 MCP host 提供统一状态资源。",
            "已实现只读和预演级能力",
            "diagnostics.py、B6",
        ],
        [
            "报告与 Office 线索",
            "清点 AFReportMSOffice、报告模板、Office proxy 标记和 GUI 报告事件，读取结果类文件、QuickLink 导出、求解器日志和报告日志，并生成轻量结果证据包计划。",
            "P0 已完成证据包基础闭环",
            "report.py、results.py、B6、B10",
        ],
        [
            "AF_API 二次开发",
            "列出 friction、heattransfer、oneelementpost 样例模块，检查 C 编译环境，生成 starter 文件计划和编译命令预览。",
            "已实现模板和预览，当前编译器缺失",
            "af_api.py、autoform_check_af_api_build_env",
        ],
        [
            "帮助主题与能力覆盖",
            "读取 helpLinks.cfg，按关键词筛选帮助主题，把 GUI 主题映射到 Agent 能力域。",
            "已实现索引级能力",
            "coverage.py、inventory.py",
        ],
        [
            "本地前端演示",
            "启动 HTTP bridge 与静态前端，展示状态、操作流、API 配置和本地只读能力。",
            "已调整为 API runtime 路线，前端通过 /api/agent 调用后端运行时",
            "frontend\\README.md、http_bridge.py",
        ],
        [
            "发布与安装检查",
            "检查 README、安装说明、卸载说明、许可、贡献说明、发布检查表、环境文件和配置模板，生成安装检查计划、源代码发布包计划和公开发布扫描结果。",
            "V1.0 已完成发布检查闭环",
            "release.py、safety.py、INSTALL.md、UNINSTALL.md、RELEASE_CHECKLIST.md、B10",
        ],
        [
            "权限、回滚与扩展边界",
            "为 ProgramData 写入目标生成备份和回滚计划，汇总已确认的 AutoForm 外部 CLI、QuickLink Export 脚本和报告模板扩展路径。",
            "P1、P2 已完成 1.0 边界说明",
            "safety.py、extension.py、B11",
        ],
    ]
    add_table(document, ["能力域", "可接收的自然语言指令范围", "成熟度", "依据"], overview_rows, [4.0, 13.5, 5.0, 4.0])

    document.add_heading("2.2 细分条目表", level=2)
    detail_rows = [
        [
            "安装与环境",
            "autoform_discover_installation、autoform_list_executables、autoform_environment_snapshot",
            "“看看这台机器有没有 AutoForm”、“当前版本是什么”、“把环境快照给我”。",
            "可直接用于汇报和诊断，输出结构化 JSON。",
        ],
        [
            "状态资源",
            "autoform://status、autoform_status_snapshot、python -m autoform_agent.cli status",
            "“先看 MCP 服务和本机 AutoForm 是否正常”、“给我一份当前状态 JSON”。",
            "本轮新增基础版；同一函数同时服务 MCP resource、MCP tool 和 CLI。",
        ],
        [
            "工程文件",
            "autoform_list_example_projects、autoform_resolve_project、autoform_project_run、autoform_example_project_baseline、autoform_inspect_afd、autoform_get_afd_readable_index、autoform_get_afd_project_summary",
            "“列出示例工程”、“这个 .afd 是什么项目”、“抽一下工程名和材料”、“把 Solver_R13 跑一遍”。",
            "当前摘要是候选字段，正式技术结论建议再用 QuickLink 导出交叉验证；工程运行链路已经能复制项目、执行 kinematic 求解并写出运行清单。",
        ],
        [
            "GUI 操作",
            "autoform_start_ui、autoform_open_afd",
            "“打开 AutoForm”、“打开 Solver_R13 工程”。",
            "MCP 默认 dry_run=True；真实 GUI 启动会产生外部进程。",
        ],
        [
            "作业与求解器",
            "autoform_forming_job_check_plan、autoform_forming_solver_kinematic_plan、autoform_forming_solver_full_plan、autoform_job_submit、autoform_job_status、autoform_job_wait、autoform_job_cancel、autoform_job_logs、autoform_job_archive、autoform_list_jobs",
            "“给这个工程做运动学检查”、“提交这个命令并登记作业”、“等这个作业结束”、“把作业日志和归档计划给我”。",
            "已有 7 个官方示例批量运动学探测证据、3 个官方示例默认完整求解返回码 0 记录，以及文件式作业生命周期测试。",
        ],
        [
            "QuickLink",
            "autoform_install_quicklink_bridge、autoform_list_quicklink_exports、autoform_parse_quicklink_xml、autoform_quicklink_schema、autoform_get_quicklink_process_plan、autoform_get_quicklink_evaluation、autoform_get_quicklink_die_face、autoform_compare_quicklink_exports",
            "“把 QuickLink 导出收集起来”、“解析 ProcessPlan”、“输出 1.0 schema”、“比较两次导出差异”。",
            "适合做 AutoForm 数据出口；V1.0 已把 ProjectData、Blank、ProcessPlan、Evaluation、DieFace 和几何文件汇总为稳定 JSON 结构。",
        ],
        [
            "材料",
            "autoform_install_materials、autoform_list_material_libraries、autoform_find_duplicate_material_files、autoform_material_library_backup_plan、autoform_inspect_material_file、autoform_mat_to_mtb_convert",
            "“安装材料库”、“查重复材料”、“把这个 .mat 转成 .mtb”。",
            "写入材料库和转换执行均需要显式确认或 execute=True。",
        ],
        [
            "队列与远程",
            "autoform_get_queue_config、autoform_get_remote_hosts、autoform_queue_health_check、autoform_queue_client_probe、autoform_lsf_command_plan",
            "“读一下队列配置”、“生成 LSF 提交命令”、“检查队列进程”。",
            "当前可读本机 Queue1 和 localhost 远程主机配置；生产集群需现场验证。",
        ],
        [
            "报告与结果线索",
            "autoform_report_inventory、autoform_report_log_events、autoform_report_ms_office_plan、autoform_postsolve_plan、autoform_result_inventory、autoform_report_delivery_plan、autoform_copy_result_evidence",
            "“查报告模板”、“从日志里找导出动作”、“清点结果证据”、“生成一个结果交付包计划”。",
            "已经能定位报告模板、Office 代理标记、结果类文件、QuickLink 导出和日志事件；真实 Office 自动导出仍需结合 AutoForm 现场权限验证。",
        ],
        [
            "AF_API",
            "autoform_list_af_api_modules、autoform_check_af_api_build_env、autoform_af_api_template_plan、autoform_af_api_build_preview",
            "“给我生成摩擦子程序模板”、“检查能不能编译用户传热模块”。",
            "当前本机 cl、icl、gcc 均未发现，AF_HOME_LIB 也未设置。",
        ],
        [
            "发布检查",
            "autoform_release_readiness_check、autoform_release_package_plan、autoform_install_check_plan、autoform_public_release_scan、autoform_write_safety_plan、autoform_internal_extension_boundary",
            "“检查 1.0 发布准备好了没有”、“生成安装检查命令”、“扫描公开发布阻断项”、“说明 AutoForm 扩展边界”。",
            "release-readiness 已返回 ready=true；public-release-scan 已返回 safe_to_publish=true；许可证为 MIT，版本号为 1.0.0。",
        ],
    ]
    add_table(document, ["细分范围", "主要 MCP 工具", "可接收表达", "说明"], detail_rows, [3.6, 8.2, 7.6, 7.0])


def add_example_project_section(document: Document) -> None:
    """Summarize how far each official example project can be driven."""
    document.add_heading("三、官方示例工程的实现程度", level=1)
    add_paragraph(
        document,
        "本节按“自然语言、脚本化、软件操作”三个层次描述 7 个官方示例工程。"
        "自然语言层指用户可直接下达的意图；脚本化层指当前 CLI 或 MCP 已能形成的结构化命令；"
        "软件操作层指已经进入 AutoForm 命令、GUI 或求解器的程度。",
    )

    rows = [
        [
            "AutoComp_R13.afd",
            "AutoForm Forming R13 AutoComp Test File；材料 DC04，厚度 1.0000 mm；候选信息显示 DieFace 未使用，TriboForm 使用。",
            "可询问工程摘要、材料、厚度、插件使用线索；可请求打开工程和运动学检查。",
            "可用 open-afd、afd-project-summary、project-run、example-baseline。",
            "已通过 MCP stdio 调用 autoform_project_run 完成 kinematic 执行，返回码 0，simulation_successful=true；V1.0 将其作为官方示例基准之一。",
        ],
        [
            "PhaseChange_R13.afd",
            "AutoForm Forming R13 PhaseChange Plug-in Test File；材料 22MnB5，厚度 1.0000 mm；DieFace 使用。",
            "可询问热成形或相变插件相关示例的工程摘要；可请求运动学检查和默认完整求解。",
            "可用 open-afd、afd-project-summary、forming-solver-kinematic-plan、forming-solver-full-plan、project-run。",
            "已纳入 docs\\example_project_baselines.json；基准记录 AF_HOME_LIB 推荐设置，便于后续现场执行。",
        ],
        [
            "Sigma_R13.afd",
            "AutoForm Forming R13 Sigma Test File；材料 DX54D，厚度 1.0000 mm；DieFace 使用。",
            "可询问 Sigma 示例的材料、特征和可读字段；可请求运动学检查。",
            "可用 open-afd、afd-project-summary、project-run、example-baseline。",
            "已纳入 docs\\example_project_baselines.json；V1.0 公开边界记录计划和候选摘要。",
        ],
        [
            "Solver_R13.afd",
            "AutoForm Forming R13 Solver Test File；材料 DC04，厚度 1.0000 mm；DieFace 使用。",
            "可作为当前主展示样例：读取摘要、打开 GUI、预演作业检查、执行运动学检查和默认完整求解。",
            "可用 open-afd、project-run、forming-job-check-plan、forming-solver-kinematic-plan、forming-solver-full-plan、solver-log-events。",
            "已通过 MCP stdio 调用 autoform_project_run 复制运行副本并执行运动学求解，返回码 0，stdout 摘要包含 simulation_successful=true。它作为 V1.0 演示主线。",
        ],
        [
            "Thermo_R13.afd",
            "AutoForm Forming R13 Thermo Plug-In Test File；材料 DX57-GI_Vegter+thermo_0.4-0.8，厚度 1.0000 mm；DieFace 使用。",
            "可询问热耦合示例摘要；可请求运动学检查。",
            "可用 open-afd、afd-project-summary、project-run、example-baseline。",
            "已纳入 docs\\example_project_baselines.json；热相关完整求解建议在现场许可证条件下继续补充执行记录。",
        ],
        [
            "Triboform_R13.afd",
            "AutoForm Forming R13 TriboForm Plug-In Test File；材料 DX54D，厚度 1.0000 mm；DieFace 使用，TriboForm 使用。",
            "可询问摩擦插件示例摘要；可请求运动学检查。",
            "可用 open-afd、afd-project-summary、project-run、example-baseline。",
            "已纳入 docs\\example_project_baselines.json；TriboForm 插件边界在公开版本中以基准和计划形式说明。",
        ],
        [
            "Trim_R13.afd",
            "AutoForm Forming R13 Trim Test File；材料 DX54D，厚度 1.0000 mm；DieFace 使用。",
            "可询问修边示例摘要；可请求运动学检查和默认完整求解。",
            "可用 open-afd、project-run、forming-solver-kinematic-plan、forming-solver-full-plan。",
            "已通过 MCP stdio 调用 autoform_project_run 完成 kinematic 执行，返回码 0，simulation_successful=true；适合作为 Solver_R13 之外的第二条演示链路。",
        ],
    ]
    add_table(document, ["示例工程", "当前可读摘要", "自然语言层", "脚本化层", "软件操作层"], rows, [3.2, 7.4, 5.5, 5.4, 5.0])

    add_paragraph(
        document,
        "总体判断：当前已经能把用户的常见汇报型指令转换为结构化 MCP 或 CLI 调用，并能先用状态资源复核本机环境，再对官方示例完成读取、打开、命令计划、运行目录生成和 Solver_R13 实际运动学求解。"
        "P0、P1、P2 已按 V1.0 公开边界收口；后续增强工作集中在更多真实结果样本、结果图像导出、企业安装包和更多 AutoForm 版本验证。",
    )


def add_open_source_comparison_section(document: Document) -> None:
    """Compare current AutoForm MCP work with Cai-aa/abaqus-mcp as a public reference."""
    document.add_heading("四、与公开 MCP 项目的差距", level=1)
    add_paragraph(
        document,
        "公开项目 Cai-aa/abaqus-mcp 可作为当前阶段的明确参照。其 README 将架构描述为外部 FastMCP server 与 Abaqus/CAE 内部插件之间的文件式 IPC，"
        "并提供模型信息读取、脚本执行、作业提交、ODB 读取、视口截图和状态资源。下表基于该 README 与本项目当前源码状态形成对比。",
    )

    rows = [
        [
            "软件内插件或常驻代理",
            "Abaqus 侧有 abaqus_mcp_plugin.py，可在 Abaqus/CAE 内运行并消费命令文件。",
            "V1.0 已把 AutoForm 自动化边界限定为外部 CLI/MCP、QuickLink Export 脚本桥接和报告模板线索，并通过 extension-boundary 输出证据。",
            "后续若获得 AutoForm 官方常驻脚本、插件、宏或事件入口依据，可在该边界上继续扩展。",
        ],
        [
            "通用脚本执行",
            "提供 execute_script，可把 Python 脚本送入 Abaqus/CAE 执行。",
            "当前无等价的 AutoForm 内部通用脚本执行能力；已有能力是受控命令计划、QuickLink 收集和专用工具。",
            "需确认 AutoForm 是否有安全可控的脚本宿主。缺少官方依据时，应保持专用工具白名单。",
        ],
        [
            "模型对象信息",
            "提供 get_model_info，README 列出 parts、materials、steps、loads、BCs、interactions、assembly instances。",
            "当前可通过 .afd 可读片段和 QuickLink archive 读取候选工程字段与导出 XML 段落，并通过 quicklink-schema 输出 V1.0 稳定字段。",
            "后续可在更多真实导出样本基础上继续扩展材料、工艺、工具、边界和评价区的字段覆盖。",
        ],
        [
            "作业生命周期",
            "提供 list_jobs 和 submit_job，并等待作业完成。",
            "当前已有 AFFormingJob 检查计划、AFFormingSolver 运动学和完整求解计划、批量探测、日志解析、作业登记工具，并新增 project-run 工程级运行清单。",
            "后续重点是把 AutoForm 队列位置、许可证等待和失败分类继续补充进同一份作业 manifest。",
        ],
        [
            "结果文件读取",
            "提供 get_odb_info，用只读方式打开 ODB 并返回 steps、frames、instances 等元数据。",
            "当前已提供 result-inventory、report-delivery-plan 和 result-evidence-copy，用只读方式汇总结果类文件、QuickLink 导出、solver 日志和报告日志，并可生成轻量证据包。",
            "后续重点是确认 AutoForm 结果对象的官方字段格式，并把图像、数值评价和 Office 报告导出纳入稳定 schema。",
        ],
        [
            "视图与图片",
            "提供 get_viewport_image，返回 Abaqus viewport screenshot 的 base64 图像。",
            "当前没有 AutoForm 视口截图工具；报告模板和 QuickLink 图像查看器线索已被 report_inventory 发现。",
            "V1.0 以结果证据包、QuickLink schema 和报告模板线索作为公开边界；结果图导出进入后续增强项。",
        ],
        [
            "实时状态资源",
            "提供 abaqus://status MCP resource，status.json 每 2 秒更新心跳。",
            "本轮新增 autoform://status MCP resource，并提供 autoform_status_snapshot 工具和 CLI status 命令；状态内容包括项目版本、端口、安装、队列、QuickLink、最近日志、覆盖矩阵和错误列表。",
            "后续应补充前端轮询展示、最近操作审计和状态历史保留策略。",
        ],
        [
            "日志与过期命令清理",
            "README 记录统一日志 mcp.log、status.json、commands、results 和 stale command cleanup。",
            "当前有 launcher 日志、诊断包计划、GUI 日志解析、作业登记目录、结果证据包和输出目录清理边界。",
            "后续可继续定义更完整的操作审计、结果归档保留周期和清理策略。",
        ],
        [
            "用户安装与菜单",
            "提供 Abaqus 插件菜单 start、stop、status，以及 .mcp.json 配置示例。",
            "当前有 start_autoform_agent.ps1、cmd 启动器、可选 MCP 配置模板、本地前端、INSTALL、UNINSTALL、install-check-plan 和 release-package-plan。",
            "V1.0 已通过环境变量覆盖项、安装说明、新手指南和 release-readiness 建立跨机器适配入口；企业内部发布渠道适配进入后续增强项。",
        ],
        [
            "开源发布成熟度",
            "GitHub API 显示 Cai-aa/abaqus-mcp 是 MIT License 公开项目，有 stars、forks、issues 和清晰 README。",
            "当前项目 pyproject 版本为 1.0.0，许可证为 MIT；已补齐 LICENSE、CONTRIBUTING、CHANGELOG、RELEASE_CHECKLIST、安装说明、卸载说明、公开发布扫描和 release-readiness 检查。",
            "公开发布操作包括推送 GitHub、设置仓库可见性为 public、创建 V1.0 标记，并保留本机测试和运行证据。",
        ],
    ]
    add_table(document, ["对比维度", "Cai-aa/abaqus-mcp", "当前 AutoForm MCP", "应补能力"], rows, [3.4, 7.4, 7.4, 8.0])


def add_next_work_section(document: Document) -> None:
    """List unfinished work in an order that supports a 1.0 release."""
    document.add_heading("五、P0、P1、P2 收口结果", level=1)
    add_paragraph(
        document,
        "本轮先完成 P0 的状态资源、作业生命周期、结果证据包和发布检查四类基础闭环，随后继续完成 P1 与 P2 的 V1.0 公开边界项。"
        "下表记录每一项的当前完成证据和后续可增强方向。",
    )
    rows = [
        [
            "P0 完成",
            "状态资源和健康检查",
            "已新增 autoform://status、autoform_status_snapshot 和 CLI status；状态快照覆盖项目版本、端口、安装、队列、QuickLink、日志、覆盖矩阵和错误列表。",
            "B3、B4、B5、B10，autoform_agent\\diagnostics.py，tests\\test_diagnostics.py。",
        ],
        [
            "P0 完成",
            "作业生命周期闭环",
            "在现有 forming_job_check_plan 和 solver probe 基础上补齐 submit、status、cancel、wait、logs、archive 和 list，默认 execute=False。",
            "autoform_agent\\jobs.py，autoform_agent\\cli.py，autoform_agent\\mcp_server.py，tests\\test_jobs.py。",
        ],
        [
            "P0 完成",
            "结果读取和报告导出",
            "以 QuickLink、AFReportMSOffice、报告模板、结果类文件和 GUI 日志为入口，形成 result-inventory、report-delivery-plan 和 result-evidence-copy。",
            "autoform_agent\\results.py，autoform_agent\\report.py，tests\\test_results.py。",
        ],
        [
            "P0 完成",
            "发布包与安装流程",
            "补齐 LICENSE、CONTRIBUTING、INSTALL、UNINSTALL、RELEASE_CHECKLIST，并提供 release-readiness、install-check-plan 和 release-package-plan。",
            "autoform_agent\\release.py，README.md，docs\\beginner_onboarding_zh.md，tests\\test_release.py。",
        ],
        [
            "P1 完成",
            "QuickLink schema 化",
            "已新增 quicklink-schema，把 ProjectData、Blank、ProcessPlan、Evaluation、DieFace、几何文件和 archive 成员数量映射为 schema_version=1.0 的稳定 JSON。",
            "autoform_agent\\quicklink.py，tests\\test_quicklink.py，B11。",
        ],
        [
            "P1 完成",
            "示例工程基准集",
            "已生成 docs\\example_project_baselines.json，记录 7 个官方示例的候选摘要、运动学求解计划和完整求解计划；Solver_R13 已完成实际 kinematic 求解。",
            "autoform_agent\\project_workflow.py，tests\\test_project_workflow.py，B10、B11。",
        ],
        [
            "P1 完成",
            "权限和回滚",
            "已新增 write-safety-plan，为 ProgramData 写入目标生成 parent_writable、backup_path 和 rollback_action；公开发布扫描已纳入 release-readiness。",
            "autoform_agent\\safety.py，tests\\test_safety_extension.py，B10、B11。",
        ],
        [
            "P1 完成",
            "前端与 API runtime 交互完善",
            "前端经 HTTP bridge 调用 agent_runtime 的职责已记录在 README、docs\\api_runtime_call_chain.md 和新手指南中；runtime 已暴露工程解析与工程运行计划工具。",
            "autoform_agent\\agent_runtime.py，frontend\\app.js，docs\\api_runtime_call_chain.md。",
        ],
        [
            "P2 完成",
            "AutoForm 内部扩展路径调研",
            "已新增 extension-boundary，记录本机确认的外部 CLI、QuickLink Export 脚本和报告模板路径，并把缺少证据的内部通用脚本宿主列为 V1.0 边界外能力。",
            "autoform_agent\\extension.py，tests\\test_safety_extension.py，B11。",
        ],
        [
            "P2 完成",
            "跨版本和跨机器适配",
            "已在 paths.py 和 .env.example 中加入 AutoForm 安装目录、ProgramData、材料目录、脚本目录、示例工程目录、QuickLink 模板、系统配置和帮助链接覆盖项。",
            "autoform_agent\\paths.py，tests\\test_paths.py，.env.example。",
        ],
    ]
    add_table(document, ["优先级", "功能项", "完成内容", "依据或原因"], rows, [1.8, 5.2, 13.2, 6.4])


def add_version_distance_section(document: Document) -> None:
    """Give a pragmatic estimate of distance to a public 1.0 release."""
    document.add_heading("六、距离可公开使用的 1.0 版本还有多远", level=1)
    add_paragraph(
        document,
        "从当前证据看，本项目已经达到可公开使用的 V1.0 本机发布标准。它有 88 个 MCP 工具入口、1 个 MCP resource、91 个 CLI 入口、81 个测试用例通过，"
        "并且已经形成状态快照、AutoForm 安装发现、材料治理、QuickLink 收集与 schema 化、示例工程读取、工程级运行、求解器计划、作业生命周期、结果证据包、公开发布扫描和发布检查。"
    )
    add_paragraph(
        document,
        "公开 V1.0 的目标已经落地为：用户能安装、能检查环境、能打开示例工程，能通过 MCP stdio 客户端完成 Solver_R13、Trim_R13 和 AutoComp_R13 三个官方示例的运动学求解、"
        "能收集 QuickLink、能执行材料治理、能生成最小结果证据包，并能在失败时获得清楚的日志和下一步动作。"
    )

    rows = [
        [
            "当前阶段",
            "V1.0 本机公开发布阶段",
            "pyproject.toml 版本为 1.0.0；81 个测试通过；AutoForm R13 本机链路清楚；P0、P1、P2 已按 V1.0 公开边界形成可复核证据。",
        ],
        [
            "1.0 核心门槛",
            "可安装、可复现、可诊断、可回滚",
            "安装说明、卸载说明、发布检查、作业登记、工程运行、结果证据包、公开发布扫描和回滚计划已经落地；许可证为 MIT，版本号为 1.0.0。",
        ],
        [
            "距离判断",
            "当前开发缺口按 V1.0 范围计为 0",
            "本判断基于 release-readiness=true、public-release-scan safe_to_publish=true、三个官方示例的 MCP project-run 执行成功、QuickLink schema 输出成功、全量测试通过和文档已同步。",
        ],
        [
            "发布建议",
            "按 V1.0 标记推送并公开仓库",
            "建议用 release-readiness、public-release-scan、install-check-plan、MCP project-run 示例组、QuickLink schema、result-inventory 和 report-delivery-plan 组成发布验收脚本。",
        ],
        [
            "后续风险",
            "更多 AutoForm 版本、企业许可证和结果视图导出仍需持续补证",
            "V1.0 采用外部命令、QuickLink 导出、报告模板线索和结果证据包作为稳定边界；更多内部对象接口进入后续版本规划。",
        ],
    ]
    add_table(document, ["判断项", "结论", "依据或说明"], rows, [3.2, 6.6, 16.6])


def add_appendix_section(document: Document) -> None:
    """Add concrete command examples and source URLs for the speaker."""
    document.add_heading("七、汇报可用命令示例与引用链接", level=1)
    command_rows = [
        ["发现安装", r"python -m autoform_agent.cli discover"],
        ["读取状态快照", r"python -m autoform_agent.cli status"],
        ["列出示例工程", r"python -m autoform_agent.cli example-projects"],
        ["读取工程摘要", r"python -m autoform_agent.cli afd-project-summary C:\ProgramData\AutoForm\AFplus\R13F\test\Solver_R13.afd"],
        ["预演打开工程", r"python -m autoform_agent.cli open-afd C:\ProgramData\AutoForm\AFplus\R13F\test\Solver_R13.afd --dry-run"],
        ["解析工程输入", r"python -m autoform_agent.cli resolve-project --example Solver_R13"],
        ["预演工程运行", r"python -m autoform_agent.cli project-run --example Solver_R13 --mode kinematic --threads 1 --output-root output\project_runs"],
        ["执行工程运行", r"python -m autoform_agent.cli project-run --example Solver_R13 --mode kinematic --threads 1 --output-root output\project_runs --execute --timeout 120"],
        ["刷新示例基准", r"python -m autoform_agent.cli example-baseline --output docs\example_project_baselines.json --threads 1"],
        ["生成运动学检查命令", r"python -m autoform_agent.cli forming-solver-kinematic-plan C:\ProgramData\AutoForm\AFplus\R13F\test\Solver_R13.afd --threads 1"],
        ["生成默认完整求解命令", r"python -m autoform_agent.cli forming-solver-full-plan C:\ProgramData\AutoForm\AFplus\R13F\test\Solver_R13.afd --threads 1"],
        ["预演登记作业", r"python -m autoform_agent.cli job-submit --name status_check -- python -m autoform_agent.cli status"],
        ["读取作业状态", r"python -m autoform_agent.cli job-status <job_id>"],
        ["读取作业日志", r"python -m autoform_agent.cli job-registered-logs <job_id>"],
        ["清点结果证据", r"python -m autoform_agent.cli result-inventory --limit 20"],
        ["预演结果证据包", r"python -m autoform_agent.cli report-delivery-plan output\result_package --limit 20"],
        ["规范化 QuickLink", r"python -m autoform_agent.cli quicklink-schema autoform_agent_data\quicklink\20260525_234139\quicklinkExport.zip"],
        ["检查发布就绪", r"python -m autoform_agent.cli release-readiness"],
        ["公开发布扫描", r"python -m autoform_agent.cli public-release-scan"],
        ["生成写入回滚计划", r"python -m autoform_agent.cli write-safety-plan C:\ProgramData\AutoForm\AFplus\R13F\scripts\CodexAgentBridge.cmd --backup-root output\rollback"],
        ["说明扩展边界", r"python -m autoform_agent.cli extension-boundary --workspace <repo-root>"],
        ["预演发布包", r"python -m autoform_agent.cli release-package-plan output\release\autoform-agent-1.0"],
        ["列出 QuickLink 导出", r"python -m autoform_agent.cli quicklink-list --workspace <repo-root>"],
        ["列出材料库", r"python -m autoform_agent.cli material-libraries"],
        ["运行测试", r"$env:TEMP='<repo-root>\tmp\pytest_runtime_final_<timestamp>'; $env:TMP=$env:TEMP; <python> -m pytest -q --basetemp tmp\pytest_basetemp_final_<timestamp>"],
    ]
    add_table(document, ["用途", "命令"], command_rows, [5.0, 21.5])

    source_rows = [
        ["本项目 README", r"README.md"],
        ["本项目开发者指南", r"DEVELOPERS.md"],
        ["MCP 工具入口", r"autoform_agent\mcp_server.py"],
        ["状态快照实现", r"autoform_agent\diagnostics.py"],
        ["作业生命周期实现", r"autoform_agent\jobs.py"],
        ["工程运行实现", r"autoform_agent\project_workflow.py"],
        ["结果证据实现", r"autoform_agent\results.py"],
        ["发布检查实现", r"autoform_agent\release.py"],
        ["公开发布扫描与回滚", r"autoform_agent\safety.py"],
        ["扩展边界实现", r"autoform_agent\extension.py"],
        ["模块覆盖矩阵", r"autoform_agent\coverage.py"],
        ["示例工程基准", r"docs\example_project_baselines.json"],
        ["示例工程摘要实现", r"autoform_agent\inventory.py"],
        ["求解器实现", r"autoform_agent\solver.py"],
        ["QuickLink 实现", r"autoform_agent\quicklink.py"],
        ["安装说明", r"INSTALL.md"],
        ["卸载说明", r"UNINSTALL.md"],
        ["发布检查表", r"RELEASE_CHECKLIST.md"],
        ["Cai-aa/abaqus-mcp README", "https://raw.githubusercontent.com/Cai-aa/abaqus-mcp/main/README.md"],
        ["Cai-aa/abaqus-mcp GitHub API", "https://api.github.com/repos/Cai-aa/abaqus-mcp"],
    ]
    add_table(document, ["引用对象", "位置"], source_rows, [6.0, 20.5])


def add_paragraph(document: Document, text: str) -> None:
    """Add a paragraph with the spacing used throughout the report."""
    paragraph = document.add_paragraph(text)
    paragraph.paragraph_format.space_after = Pt(5)
    paragraph.paragraph_format.line_spacing = 1.15


def add_table(document: Document, headers: list[str], rows: list[list[str]], widths_cm: list[float]) -> None:
    """Create a styled table with fixed column widths for consistent Word layout."""
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False

    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = header
        set_cell_width(cell, widths_cm[index])
        shade_cell(cell, "D9EAF7")
        make_cell_bold(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = value
            set_cell_width(cells[index], widths_cm[index])
            cells[index].vertical_alignment = WD_ALIGN_VERTICAL.TOP
            for paragraph in cells[index].paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.05
                for run in paragraph.runs:
                    run.font.size = Pt(8)
                    if looks_like_code(value):
                        run.font.name = "Consolas"
                        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    document.add_paragraph()


def set_cell_width(cell, width_cm: float) -> None:
    """Write an explicit Word table-cell width in twentieths of a point."""
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_cm * 567)))
    tc_w.set(qn("w:type"), "dxa")


def shade_cell(cell, fill: str) -> None:
    """Apply a simple background fill to a header cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def make_cell_bold(cell) -> None:
    """Make every run in a header cell bold and dark."""
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x1F, 0x1F, 0x1F)


def looks_like_code(value: str) -> bool:
    """Detect cells that contain command text or paths and benefit from a monospace font."""
    markers = ["\\", ".py", ".afd", ".exe", ".cmd", "python -m", "--", "autoform_", "http"]
    return any(marker in value for marker in markers)


if __name__ == "__main__":
    main()
