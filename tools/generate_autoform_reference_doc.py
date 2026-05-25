from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


WORKSPACE = Path(r"F:\【项目和任务】\EIT\2026\AUTO_AutoForm")
OUTPUT_PATH = WORKSPACE / "output" / "doc" / "AutoForm官方命令与Agent指令对照表.docx"


def main() -> None:
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    configure_document(document)

    add_title(document)
    add_scope(document)
    add_source_table(document)
    add_official_command_tables(document)
    add_agent_tables(document)
    add_mapping_table(document)
    add_usage_notes(document)
    add_appendix(document)

    document.save(OUTPUT_PATH)
    print(OUTPUT_PATH)


def configure_document(document: Document) -> None:
    """Apply a compact landscape layout suited for wide command tables."""
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Cm(1.4)
    section.bottom_margin = Cm(1.4)
    section.left_margin = Cm(1.2)
    section.right_margin = Cm(1.2)

    normal = document.styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(9)

    for style_name, size in [("Title", 18), ("Heading 1", 13), ("Heading 2", 11)]:
        style = document.styles[style_name]
        style.font.name = "宋体"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        style.font.size = Pt(size)
        style.font.bold = True


def add_title(document: Document) -> None:
    title = document.add_heading("AutoForm 官方命令与 Agent 指令对照表", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("适用对象：本机 AutoForm Forming R13 与当前 autoform_agent 项目").bold = True
    document.add_paragraph(
        "生成日期：2026-05-22。本文档依据本机已安装 AutoForm 文件、Windows 注册表、开始菜单快捷方式、"
        "AutoForm ProgramData 脚本目录、AF_API 头文件，以及当前工作区源码整理。"
    )


def add_scope(document: Document) -> None:
    document.add_heading("一、范围与判定原则", level=1)
    rows = [
        [
            "官方条目收录原则",
            "只收录本轮能从本机 AutoForm 安装目录、ProgramData、注册表、快捷方式或随安装包提供的头文件中直接验证的命令、脚本指令、配置项和 API 例程。"
        ],
        [
            "发现的 exe",
            "bin 目录中的 exe 会单独列出为“可执行入口清单”。未找到随附参数说明的 exe，只记录其存在和可能用途，不扩展为已确认命令语法。"
        ],
        [
            "Agent 条目收录原则",
            "以当前工作区源码为准，包含 CLI 子命令、MCP 工具、关键参数、默认行为、底层调用关系和安全提示。"
        ],
        [
            "执行安全",
            "涉及启动 GUI、写入材料库、安装 QuickLink 脚本、清理队列文件和提交计算任务的条目，均建议先使用 dry run 或只读验证。"
        ],
    ]
    add_table(document, ["项目", "说明"], rows, [4.0, 22.5])


def add_source_table(document: Document) -> None:
    document.add_heading("二、资料来源清单", level=1)
    rows = [
        ["S1", r"Windows 卸载注册表", "AutoForm Forming R13，版本 13.0.1.02，安装日期 20260519，安装位置 D:\\Program Files\\AutoForm\\AFplus\\R13F。"],
        ["S2", r"D:\Program Files\AutoForm\AFplus\R13F\package_info_lite.json", "产品为 Forming，系统为 Windows，分支 AFProducts/R13/R13.0.1.2，构建时间戳 20251014161456。"],
        ["S3", r"C:\ProgramData\Microsoft\Windows\Start Menu\Programs\AutoForm\Forming R13\*.lnk", "开始菜单快捷方式指向 AFSplash.exe 与 AFHostInfo.exe。Forming 快捷方式参数为 -language=en -afformingui。"],
        ["S4", r"HKCR\AFForming.Design\shell\open\command", r".afd 文件关联命令为 AFFormingUI.exe -file %%1。"],
        ["S5", r"D:\Program Files\AutoForm\AFplus\R13F\bin\*.cmd", "包含 AFFormingJob_R13.cmd、AFFileServer.cmd、AFRemoteUser.cmd、AFEmailNotification.cmd、aflsf_copy.cmd、aflsf_share.cmd、killQueueServer.cmd。"],
        ["S6", r"C:\ProgramData\AutoForm\AFplus\R13F\scripts\Scripts_Readme.txt:15-24", "说明 QuickLink Export 脚本的放置目录、显示方式、执行方式以及传入参数。"],
        ["S7", r"C:\ProgramData\AutoForm\AFplus\R13F\scripts\sendMail.bat:6-18", "示例脚本展示 #EXPORTSTANDARD、#CLEANUPDATA、#SAVEDESIGN，以及 %1 接收附件路径。"],
        ["S8", r"C:\ProgramData\AutoForm\AFplus\R13F\systemConfigFile.xml", "队列、远程计算、RequestServer、日志和更新配置项。"],
        ["S9", r"D:\Program Files\AutoForm\AFplus\R13F\AF_API\af_friction.h", "用户摩擦模块的编译命令、启用控制变量和导出函数声明。"],
        ["S10", r"D:\Program Files\AutoForm\AFplus\R13F\AF_API\af_heattransfer.h", "用户传热模块的编译命令、启用控制变量和导出函数声明。"],
        ["S11", r"D:\Program Files\AutoForm\AFplus\R13F\AF_API\af_oneelementpost.h", "用户后处理模块的初始化、结束和逐增量后处理函数声明。"],
        ["S12", r"F:\【项目和任务】\EIT\2026\AUTO_AutoForm\autoform_agent\*.py", "当前 Agent 的 CLI、MCP、AutoForm 路径发现、进程调用、材料库和 QuickLink 桥接实现。"],
    ]
    add_table(document, ["编号", "来源位置", "可验证内容"], rows, [1.8, 10.0, 14.7])


def add_official_command_tables(document: Document) -> None:
    document.add_heading("三、AutoForm 官方可验证命令与指令", level=1)

    document.add_heading("3.1 启动、文件关联与批处理入口", level=2)
    rows = [
        [
            "OFF-GUI-01",
            r"AFSplash.exe -language=en -afformingui",
            "通过开始菜单启动 AutoForm Forming 图形界面。工作目录为安装目录 bin。",
            "普通用户可执行；适合用户手动启动或 Agent 通过外部进程启动。",
            "S3",
        ],
        [
            "OFF-GUI-02",
            r"AFHostInfo.exe",
            "打开 AutoForm Forming R13 HostInfo 工具，用于查看主机与环境信息。",
            "普通用户可执行；常用于排查环境、主机信息和许可相关问题。",
            "S3",
        ],
        [
            "OFF-GUI-03",
            r"AFFormingUI.exe -file <afd路径>",
            "由 Windows .afd 文件关联提供的打开工程命令。",
            "路径应使用完整路径；命令来源为 AFForming.Design 的 open command。",
            "S4",
        ],
        [
            "OFF-JOB-01",
            r"AFFormingJob_R13.cmd <原始参数>",
            "批处理包装器先把 AutoForm bin 加入 PATH，然后把所有参数传给 AFFormingJob。",
            r"AFFormingJob_R13.cmd 第 3 行设置 PATH，第 5 行执行 AFFormingJob %*。",
            "S5",
        ],
        [
            "OFF-JOB-02",
            r"AFFormingJob.exe <原始参数>",
            "实际计算作业入口。当前本机资料确认可由包装器调用，具体业务参数需依据 AutoForm 作业文件和官方作业说明补充。",
            "Agent 的 run-job 对该 exe 进行透明转发，便于后续接入已知参数。",
            "S5、S12",
        ],
    ]
    add_table(document, ["编号", "命令或指令", "解释", "使用要点", "依据"], rows, [2.2, 6.8, 8.0, 7.5, 2.0])

    document.add_heading("3.2 QuickLink Export 脚本机制与脚本内指令", level=2)
    rows = [
        [
            "OFF-QL-01",
            r"将 bat/cmd/exe 等脚本放入 C:\ProgramData\AutoForm\AFplus\R13F\scripts",
            "AutoForm Forming 会把 scripts 目录中的脚本加入 QuickLink Export scripts 选项。",
            "txt 文件不作为脚本入口加入；脚本名就是用户在 QuickLink Export 中选择的入口。",
            "S6",
        ],
        [
            "OFF-QL-02",
            r"<script> %1",
            "AutoForm 调用 QuickLink 脚本时，会在系统默认临时目录创建 QuickLink transfer package，并把 archive 路径作为第一个参数传给脚本。",
            "脚本需要把该临时数据转移到用户定义的位置；脚本结束后 AutoForm 会清理临时位置。",
            "S6、S7",
        ],
        [
            "OFF-QL-03",
            r"#EXPORTSTANDARD <标准xml>",
            "QuickLink 示例脚本中的元指令，用于指定 QuickLinkExportStandard。",
            "sendMail.bat 示例给出占位文件名 UseNameOfQuickLinkExportStandard.xml，实际项目需替换为标准目录中存在的标准文件。",
            "S7",
        ],
        [
            "OFF-QL-04",
            r"#CLEANUPDATA",
            "QuickLink 示例脚本中的元指令，表示导出流程包含数据清理行为。",
            "随 AutoForm 示例脚本提供；具体清理范围应以对应版本 QuickLink 行为为准。",
            "S7",
        ],
        [
            "OFF-QL-05",
            r"#SAVEDESIGN",
            "QuickLink 示例脚本中的元指令，表示导出流程会保存设计。",
            "适合需要把当前设计状态随 QuickLink 包一起传递的脚本。",
            "S7",
        ],
        [
            "OFF-QL-06",
            r"set attachment=%1",
            "sendMail.bat 示例中把 AutoForm 传入的 QuickLink archive 路径作为邮件附件路径。",
            "后续使用 Outlook /a 参数把该路径附加到新邮件。",
            "S7",
        ],
    ]
    add_table(document, ["编号", "命令或指令", "解释", "使用要点", "依据"], rows, [2.2, 7.4, 7.6, 7.0, 2.3])

    document.add_heading("3.3 队列、远程计算、邮件通知与作业调度脚本", level=2)
    rows = [
        [
            "OFF-QUEUE-01",
            r"AFQueueClient -fspw",
            "AFFileServer.cmd 调用的队列客户端参数。",
            "用于文件服务相关的队列口令或服务流程，具体交互由 AutoForm 队列组件处理。",
            "S5",
        ],
        [
            "OFF-QUEUE-02",
            r"AFQueueClient -rupw",
            "AFRemoteUser.cmd 调用的队列客户端参数。",
            "用于远程用户相关的队列口令或服务流程，执行后脚本 pause 保持窗口。",
            "S5",
        ],
        [
            "OFF-QUEUE-03",
            r"killQueueServer.cmd",
            "强制结束 AFQueueServer.exe，并删除系统配置目录下队列数据文件。",
            r"包含 taskkill /f /t /im AFQueueServer.exe 和 del *.que *.pid *.lck。属于维护操作，执行前应确认没有运行中的作业。",
            "S5",
        ],
        [
            "OFF-MAIL-01",
            r"AFEmailNotification.cmd <address> <server> <text>",
            "将第 1、2、3 个参数分别作为邮件地址、SMTP 服务器和消息文本，调用 PowerShell Send-MailMessage。",
            "依赖本机 PowerShell 和 SMTP 配置；Send-MailMessage 在较新 PowerShell 中可能有弃用风险，生产环境需测试。",
            "S5",
        ],
        [
            "OFF-LSF-01",
            r"aflsf_share.cmd -bsub <commandline> <username> <jobname> <puse> <lictype> <nlics> <thermo> <workdir>",
            "LSF 共享目录模式提交入口。脚本进入工作目录，调用 bsub.exe，附加作业名和许可服务器参数。",
            "脚本内许可服务器默认 2375@chlicenseserver；需要按现场集群配置调整。",
            "S5",
        ],
        [
            "OFF-LSF-02",
            r"aflsf_share.cmd -bjobs <jobid>",
            "LSF 作业查询入口，调用 bjobs.exe 并输出结果。",
            "结果通过临时 stdout 文件转出。",
            "S5",
        ],
        [
            "OFF-LSF-03",
            r"aflsf_share.cmd -bkill <jobid>",
            "LSF 作业终止入口，调用 bkill.exe。",
            "用于取消已提交作业，建议记录 jobid 和执行人。",
            "S5",
        ],
        [
            "OFF-LSF-04",
            r"aflsf_copy.cmd -bsub ... <ninputfiles> <输入文件...> <noutputfiles> <输出文件...>",
            "LSF 拷贝模式提交入口。在 bsub 命令中追加 -f 文件传输规则，并生成 command.cmd 执行。",
            "适合需要显式声明输入和输出文件传输的集群作业。",
            "S5",
        ],
        [
            "OFF-LSF-05",
            r"aflsf_copy.cmd -bjobs <jobid> / -bkill <jobid>",
            "拷贝模式同样提供查询和终止入口，底层分别调用 bjobs.exe 与 bkill.exe。",
            "与 share 模式一致，差异主要在 -bsub 阶段是否生成文件传输规则。",
            "S5",
        ],
    ]
    add_table(document, ["编号", "命令或指令", "解释", "使用要点", "依据"], rows, [2.2, 8.6, 7.0, 6.5, 2.2])

    document.add_heading("3.4 AF_API 用户扩展接口", level=2)
    rows = [
        [
            "OFF-API-01",
            r"gcc -fPIC -c af_friction.c -o af_friction.o",
            "Linux 下编译用户摩擦模块对象文件。",
            "来自 af_friction.h 示例。后续需生成共享库 libafuser.so。",
            "S9",
        ],
        [
            "OFF-API-02",
            r"gcc -shared -o libafuser.so af_friction.o",
            "Linux 下生成用户摩擦共享库。",
            "AutoForm 会先在仿真文件所在目录查找，再查找安装目录 lib；AF_HOME_LIB 指向该目录。",
            "S9",
        ],
        [
            "OFF-API-03",
            r"icl -DWIN32 /LD /Felibafuser.dll af_friction.c 或 cl -DWIN32 /LD /Felibafuser.dll af_friction.c",
            "Windows 下生成用户摩擦动态库。",
            "启用方式为把控制变量 UserFriction 设置为 ON。",
            "S9",
        ],
        [
            "OFF-API-04",
            r"af_FrictionDataSize / af_InitFriction / af_InterpolateFrictionOnSide / af_InterpolateFrictionInElement / af_Friction",
            "用户摩擦模块函数集合，覆盖数据大小、初始化、网格细化插值和摩擦计算。",
            "核心计算函数 af_Friction 在材料流动求解开始时针对接触节点调用。",
            "S9",
        ],
        [
            "OFF-API-05",
            r"gcc -fPIC -c af_heattransfer.c -o af_heattransfer.o",
            "Linux 下编译用户传热模块对象文件。",
            "来自 af_heattransfer.h 示例。后续需生成共享库 libafuser.so。",
            "S10",
        ],
        [
            "OFF-API-06",
            r"gcc -shared -o libafuser.so af_heattransfer.o",
            "Linux 下生成用户传热共享库。",
            "AutoForm 会先在仿真文件所在目录查找，再查找安装目录 lib；AF_HOME_LIB 指向该目录。",
            "S10",
        ],
        [
            "OFF-API-07",
            r"icl -DWIN32 /LD /Felibafuser.dll af_heattransfer.c 或 cl -DWIN32 /LD /Felibafuser.dll af_heattransfer.c",
            "Windows 下生成用户传热动态库。",
            "启用方式为把控制变量 UserHeatTransfer 设置为 ON。",
            "S10",
        ],
        [
            "OFF-API-08",
            r"af_HeatTransferDataSize / af_InitHeatTransfer / af_InterpolateHeatTransferOnSide / af_InterpolateHeatTransferInElement / af_HeatTransfer",
            "用户传热模块函数集合，覆盖数据大小、初始化、插值和有效传热系数计算。",
            "af_HeatTransfer 在热求解阶段调用，输出有效传热系数。",
            "S10",
        ],
        [
            "OFF-API-09",
            r"af_OneElementInitPost / af_OneElementPost / af_OneElementFinishPost",
            "用户后处理模块接口，分别用于初始化数据、每个增量结束时执行后处理、最终释放数据。",
            "来自 af_oneelementpost.h 的导出函数声明。",
            "S11",
        ],
    ]
    add_table(document, ["编号", "命令或指令", "解释", "使用要点", "依据"], rows, [2.2, 9.4, 6.8, 6.0, 2.0])

    document.add_heading("3.5 systemConfigFile.xml 关键配置指令", level=2)
    rows = [
        ["OFF-CONF-01", "MaxJobs", "队列允许的并发作业数量。当前值为 1。", "提高并发前需结合许可证、CPU 和磁盘资源评估。", "S8"],
        ["OFF-CONF-02", "LicenseServer", "队列使用的许可服务器地址。当前值为 2375@localhost。", "远程或集中许可环境需改为现场许可服务器。", "S8"],
        ["OFF-CONF-03", "RestrictToParallelSolver", "限制并行求解器选项。注释中给出可用值 0、1、2、3、4、8、16。", "用于约束作业使用的并行能力。", "S8"],
        ["OFF-CONF-04", "RestrictQueuingOptions", "限制队列操作选项。当前值为 QueueToTop。", "注释列出 QueueToTop、QueueToMyTop、QueueToEnd。", "S8"],
        ["OFF-CONF-05", "RemoteComputingConfiguration/host/port", "AFRemoteService 端口。当前值为 865。", "端口冲突或防火墙策略需同步处理。", "S8"],
        ["OFF-CONF-06", "module/item", "远程主机支持的模块。当前包含 Sigma、Trim、Solver、Compensation。", "用于声明该主机可承接的计算模块。", "S8"],
        ["OFF-CONF-07", "listenPort / portRange", "RequestServer 监听端口和端口范围。当前 listenPort 为 0，portRange 为 50。", "listenPort 为 0 时通常由系统或程序选择可用端口。", "S8"],
        ["OFF-CONF-08", "logLevel", "日志级别。当前值为 info。注释列出 off、fatal、error、warning、info、debug、trace。", "排查问题时可临时提高到 debug 或 trace。", "S8"],
        ["OFF-CONF-09", "automaticLogCollectionLevel", "作业结束时自动收集日志的策略。当前值为 Failed。", "注释列出 Never、Failed、Always。", "S8"],
        ["OFF-CONF-10", "automaticDataCollection", "是否额外收集 afd 和 tmp 数据。当前值为 false。", "涉及数据体积和保密边界，启用前需确认存储位置。", "S8"],
        ["OFF-CONF-11", "collectedLogsLocation", "收集日志的存放位置。当前值为 AFD。", "注释说明可用 TMP、LOG、AFD 或文件路径。", "S8"],
    ]
    add_table(document, ["编号", "配置项", "解释", "使用要点", "依据"], rows, [2.2, 6.5, 7.5, 7.0, 2.0])

    document.add_heading("3.6 bin 目录可执行入口清单", level=2)
    rows = [[name, explain_executable(name)] for name in EXECUTABLES]
    add_table(document, ["可执行文件", "依据和备注"], rows, [7.0, 19.0])


def add_agent_tables(document: Document) -> None:
    document.add_heading("四、当前 Agent 指令与解释", level=1)

    document.add_heading("4.1 CLI 子命令", level=2)
    rows = [
        [
            "discover",
            r"python -m autoform_agent.cli discover",
            "发现本机 AutoForm Forming 安装，输出安装位置、bin、材料库、脚本目录、测试工程目录和 package_info。",
            "只读操作。",
            r"cli.py:17、paths.py:91-177",
        ],
        [
            "archive-list",
            r"python -m autoform_agent.cli archive-list <archive> [--limit N]",
            "列出 rar、zip、7z、tar、gz、tgz 等材料包内容，底层使用 bsdtar。",
            "只读操作；适合先查看材料包目录结构。",
            r"cli.py:19-21、materials.py:16、128-164",
        ],
        [
            "start-ui",
            r"python -m autoform_agent.cli start-ui [--graphics directx11|opengl2] [--dry-run]",
            "启动 AutoForm Forming GUI。当前实现调用 AFSplash.exe，并追加 -afformingui 和图形参数。",
            "建议先 --dry-run 查看命令。图形参数为 Agent 实现项，需结合目标 AutoForm 版本启动验证。",
            r"cli.py:23-25、process.py:10-24",
        ],
        [
            "open-afd",
            r"python -m autoform_agent.cli open-afd <afd路径> [--dry-run]",
            "打开指定 .afd 工程。当前实现调用 AFFormingUI.exe -file <afd路径>。",
            "会检查文件是否存在；GUI 启动后命令立即返回。",
            r"cli.py:27-29、process.py:27-43",
        ],
        [
            "run-job",
            r"python -m autoform_agent.cli run-job [--dry-run] [--timeout 秒] -- <AFFormingJob参数>",
            "把后续参数原样转发给 AFFormingJob.exe，用于批处理或计算作业。",
            "具体作业参数需来自 AutoForm 项目或官方作业说明；--dry-run 只打印命令。",
            r"cli.py:31-34、process.py:46-64",
        ],
        [
            "install-materials",
            r"python -m autoform_agent.cli install-materials <源目录或压缩包> [--library-name 名称] [--target-dir 目录] [--include-docs] [--dry-run] [--json]",
            "从目录或压缩包中筛选 .mat、.mtb、.csv，可选文档文件，并复制到 AutoForm 材料库目录。",
            "默认目标为 C:\\ProgramData\\AutoForm\\AFplus\\R13F\\materials 下的库目录；写入 ProgramData 可能需要管理员权限。",
            r"cli.py:36-42、materials.py:12-16、55-125",
        ],
        [
            "install-quicklink-bridge",
            r"python -m autoform_agent.cli install-quicklink-bridge [--workspace 目录] [--script-name 名称] [--dry-run]",
            "向 AutoForm scripts 目录安装 QuickLink 桥接 cmd，使 AutoForm 传入的 QuickLink archive 被当前 Agent 收集。",
            "默认脚本名 CodexAgentBridge.cmd；写入 ProgramData scripts 目录可能需要管理员权限。",
            r"cli.py:44-47、quicklink.py:9-43",
        ],
    ]
    add_table(document, ["子命令", "用法", "解释", "安全与行为", "源码依据"], rows, [2.8, 9.0, 6.5, 6.3, 2.8])

    document.add_heading("4.2 Agent 参数与文件类型规则", level=2)
    rows = [
        ["--dry-run", "适用于 start-ui、open-afd、run-job、install-materials、install-quicklink-bridge。", "只返回将要执行的命令或目标路径，不真正启动或写入。", "cli.py、process.py、materials.py、quicklink.py"],
        ["--graphics", "directx11 或 opengl2。", "映射为 -directx11 或 -opengl2。", "cli.py:24、process.py:66-72"],
        ["--limit", "archive-list 的最大输出条数。", "便于查看大压缩包时先取前 N 个条目。", "cli.py:21、56-62"],
        ["--timeout", "run-job 的最长等待秒数。", "传给 subprocess.run 的 timeout。", "cli.py:34、process.py:46-64"],
        ["--include-docs", "install-materials 可选参数。", "默认筛选材料和 csv；启用后增加 pdf、doc、docx、xls、xlsx、txt。", "materials.py:12-16、65-69"],
        ["--json", "install-materials 可选参数。", "输出完整 planned_files 与 copied_files，适合给上层 Agent 或日志系统读取。", "cli.py:42、86-99"],
        [".mat / .mtb", "AutoForm 材料文件。", "归类为 material。", "materials.py:12、77-78"],
        [".csv", "支撑数据文件。", "归类为 support。", "materials.py:14、79-80"],
        [".rar / .zip / .7z / .tar / .gz / .tgz", "可作为材料包输入。", "归类为 archive source，当前读取列表依赖 bsdtar。", "materials.py:16、99、128-164"],
    ]
    add_table(document, ["参数或类型", "适用范围", "解释", "源码依据"], rows, [4.2, 8.0, 10.0, 4.5])

    document.add_heading("4.3 MCP 工具", level=2)
    rows = [
        ["autoform_discover_installation()", "无参数", "返回发现到的 AutoForm 安装和关键路径。", "默认只读。", "mcp_server.py:21-24"],
        ["autoform_start_ui(graphics='directx11', dry_run=True)", "graphics、dry_run", "启动 AutoForm Forming 或返回启动命令。", "MCP 默认 dry_run=True，避免模型误启动 GUI。", "mcp_server.py:27-30"],
        ["autoform_open_afd(afd_path, dry_run=True)", "afd_path、dry_run", "打开 .afd 工程或返回打开命令。", "会检查工程文件存在性。", "mcp_server.py:33-36"],
        ["autoform_install_materials(source, library_name=None, include_docs=False, dry_run=True)", "source、library_name、include_docs、dry_run", "安装材料库或返回安装计划。", "MCP 默认 dry_run=True，真实写入需显式关闭。", "mcp_server.py:39-53"],
        ["autoform_install_quicklink_bridge(workspace, script_name='CodexAgentBridge.cmd', dry_run=True)", "workspace、script_name、dry_run", "安装 QuickLink 桥接脚本或返回目标路径。", "MCP 默认 dry_run=True，真实写入需显式关闭。", "mcp_server.py:56-69"],
    ]
    add_table(document, ["MCP 工具", "参数", "解释", "安全默认值", "源码依据"], rows, [6.5, 5.8, 7.2, 5.2, 2.6])

    document.add_heading("4.4 QuickLink 桥接脚本", level=2)
    rows = [
        [
            "CodexAgentBridge.cmd",
            r"set QUICKLINK_ARCHIVE=%~1",
            "读取 AutoForm QuickLink Export 传入的 archive 路径。",
            "quicklink.py:19",
        ],
        [
            "CodexAgentBridge.cmd",
            r"set PYTHONPATH=<workspace>;%PYTHONPATH%",
            "把当前工作区加入 Python 模块搜索路径，保证 AutoForm 调起脚本时能找到 autoform_agent。",
            "quicklink.py:22",
        ],
        [
            "CodexAgentBridge.cmd",
            r"<python> -m autoform_agent.quicklink_bridge %QUICKLINK_ARCHIVE% --workspace <workspace>",
            "把 QuickLink archive 交给 Python 桥接模块处理，并把结果保存到工作区。",
            "quicklink.py:23、quicklink_bridge.py",
        ],
        [
            "quicklink_bridge",
            r"autoform_agent_data\quicklink\<时间戳>\manifest.json",
            "记录 archive 原路径、收集时间、文件列表和目标目录。",
            "quicklink_bridge.py",
        ],
    ]
    add_table(document, ["组件", "指令", "解释", "源码依据"], rows, [4.0, 9.2, 10.0, 3.0])


def add_mapping_table(document: Document) -> None:
    document.add_heading("五、自然语言操作与命令映射", level=1)
    rows = [
        ["“找到我电脑上的 AutoForm”", "autoform_discover_installation 或 discover", "注册表、已知安装目录、package_info", "只读。"],
        ["“启动 AutoForm”", "autoform_start_ui 或 start-ui", "AFSplash.exe -afformingui", "先 dry run；真实启动会打开 GUI。"],
        ["“打开这个 afd 工程”", "autoform_open_afd 或 open-afd", "AFFormingUI.exe -file <afd路径>", "要求 afd 文件存在。"],
        ["“查看材料包里面有什么”", "archive-list", "bsdtar 列出压缩包成员", "只读；适合 rar 材料包预检。"],
        ["“把材料库安装进去”", "autoform_install_materials 或 install-materials", "复制 .mat、.mtb、.csv 到 AutoForm materials", "先 dry run；写 ProgramData 可能需要管理员权限。"],
        ["“把 QuickLink 导出的包交给 Agent”", "install-quicklink-bridge", "AutoForm QuickLink scripts 机制", "安装脚本后在 AutoForm QuickLink Export 中选择脚本名。"],
        ["“按已知参数跑作业”", "run-job", "AFFormingJob.exe", "参数需先由官方作业说明或样例工程确认。"],
        ["“接入集群队列”", "后续可扩展新 MCP 工具", "aflsf_share.cmd、aflsf_copy.cmd、systemConfigFile.xml", "需要现场 LSF、许可服务器和共享目录策略。"],
        ["“扩展摩擦或传热模型”", "后续可生成 C 模板和编译脚本", "AF_API 头文件和 UserFriction/UserHeatTransfer 控制变量", "需要 C 编译器和 AutoForm 模型启用控制变量。"],
    ]
    add_table(document, ["用户意图", "Agent 指令", "底层 AutoForm 依据", "注意事项"], rows, [6.0, 6.5, 8.5, 5.5])


def add_usage_notes(document: Document) -> None:
    document.add_heading("六、实施建议", level=1)
    rows = [
        ["第一阶段", "优先使用 CLI 与 MCP 的 dry run，确认路径、命令和材料库目标目录。", "风险最低，适合把自然语言任务转成可审计命令。"],
        ["第二阶段", "把 QuickLink 桥接作为 AutoForm 内部到 Agent 的数据出口。", "AutoForm 官方脚本机制已确认会把 archive 路径作为参数传入。"],
        ["第三阶段", "对 start-ui、open-afd、install-materials 开启真实执行。", "启动 GUI 和写入材料库都属于可见操作，应留日志。"],
        ["第四阶段", "在拿到 AFFormingJob 参数说明或样例作业后，封装更细的 run-job 子命令和 MCP 工具。", "避免把未知参数交给上层模型自由拼接。"],
        ["第五阶段", "面向 AF_API 生成 C 模板、编译脚本和启用检查清单。", "适合摩擦、传热和后处理二次开发。"],
    ]
    add_table(document, ["阶段", "建议动作", "理由"], rows, [3.0, 13.5, 10.0])


def add_appendix(document: Document) -> None:
    document.add_heading("七、命令示例附录", level=1)
    rows = [
        ["发现安装", r"python -m autoform_agent.cli discover"],
        ["预演启动 GUI", r"python -m autoform_agent.cli start-ui --graphics directx11 --dry-run"],
        ["预演打开工程", r"python -m autoform_agent.cli open-afd C:\ProgramData\AutoForm\AFplus\R13F\test\Solver_R13.afd --dry-run"],
        ["查看材料包", r"python -m autoform_agent.cli archive-list C:\Users\Tang Xufeng\Desktop\主机厂材料库.rar --limit 50"],
        ["预演安装材料库", r"python -m autoform_agent.cli install-materials C:\Users\Tang Xufeng\Desktop\主机厂材料库.rar --library-name 主机厂材料库 --dry-run"],
        ["输出完整安装计划", r"python -m autoform_agent.cli install-materials C:\Users\Tang Xufeng\Desktop\主机厂材料库.rar --library-name 主机厂材料库 --dry-run --json"],
        ["预演安装 QuickLink 桥接", r"python -m autoform_agent.cli install-quicklink-bridge --workspace F:\【项目和任务】\EIT\2026\AUTO_AutoForm --dry-run"],
        ["启动 MCP 服务", r"python -m autoform_agent.mcp_server"],
    ]
    add_table(document, ["用途", "命令"], rows, [5.0, 21.5])


def add_table(document: Document, headers: list[str], rows: list[list[str]], widths_cm: list[float]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False

    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = header
        set_cell_width(cell, widths_cm[index])
        shade_cell(cell, "D9EAF7")
        make_cell_bold(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = value
            set_cell_width(cells[index], widths_cm[index])
            cells[index].vertical_alignment = WD_ALIGN_VERTICAL.TOP
            for paragraph in cells[index].paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.size = Pt(8)
                    if looks_like_code(value):
                        run.font.name = "Consolas"
                        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    document.add_paragraph()


def set_cell_width(cell, width_cm: float) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_cm * 567)))
    tc_w.set(qn("w:type"), "dxa")


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def make_cell_bold(cell) -> None:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x1F, 0x1F, 0x1F)


def looks_like_code(value: str) -> bool:
    markers = ["\\", ".exe", ".cmd", "python -m", "--", "<", ">", "%", "AF", "autoform_"]
    return any(marker in value for marker in markers)


def explain_executable(name: str) -> str:
    known = {
        "AFSplash.exe": "开始菜单快捷方式确认用于启动 AutoForm Forming。",
        "AFHostInfo.exe": "开始菜单快捷方式确认用于打开主机信息工具。",
        "AFFormingUI.exe": ".afd 文件关联确认可使用 -file 打开工程。",
        "AFFormingJob.exe": "AFFormingJob_R13.cmd 确认会转发参数给该作业入口。",
        "AFQueueClient.exe": "AFFileServer.cmd 与 AFRemoteUser.cmd 确认调用该队列客户端。",
        "AFQueueServer.exe": "killQueueServer.cmd 确认该队列服务进程名。",
        "7za.exe": "压缩工具入口，存在于 AutoForm bin 目录。",
    }
    return known.get(name, "本轮确认该 exe 存在于 AutoForm bin 目录；尚未在本机随附资料中定位到参数说明，使用前需补充对应官方说明或样例。")


EXECUTABLES = [
    "7za.exe",
    "AFCloudTool.exe",
    "AFCloudToolTask.exe",
    "AFCloudToolUI.exe",
    "AFEncryption.exe",
    "afexchange.exe",
    "AFFormingCopy.exe",
    "AFFormingJob.exe",
    "AFFormingPostSolve.exe",
    "AFFormingRGen.exe",
    "AFFormingSolver.exe",
    "AFFormingUI.exe",
    "AFGuideLineConverter.exe",
    "AFHostInfo.exe",
    "AFJobStarter.exe",
    "AFMat2Mtb.exe",
    "AFOSSolver.exe",
    "AFPressConverter.exe",
    "AFProgressWindow.exe",
    "AFProtectLubrication.exe",
    "AFQueueClient.exe",
    "AFQueueServer.exe",
    "AFReconnect.exe",
    "AFRemoteService.exe",
    "AFReportMSOffice.exe",
    "AFSplash.exe",
    "AFToolShopEditor.exe",
    "AnimationWriter.exe",
]


if __name__ == "__main__":
    main()
